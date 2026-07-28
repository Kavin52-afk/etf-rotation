"""Parser for source strategy DOCX daily reports."""

from __future__ import annotations

import json
import re
import shutil
from pathlib import Path
from typing import Iterable

import pandas as pd

from .config import ProjectConfig
from .universe import ETF, etf_maps, load_universe, normalize_name
from .utils import ensure_dir

DATE_PATTERN = re.compile(r"今天是\s*(?P<date>\d{4}-\d{2}-\d{2})")
BROAD_START_PATTERN = re.compile(r"大类\s*ETF\s*截至昨日强弱排序\s*[:：]?", re.S)
SECTOR_START_PATTERN = re.compile(r"行业\s*ETF\s*截至昨日强弱排序\s*[:：]?", re.S)
YESTERDAY_PATTERN = re.compile(r"昨日收盘组合持仓\s*(?P<holdings>\[[^\]]*\])", re.S)
TARGET_PATTERN = re.compile(r"今日收盘组合应持仓\s*(?P<holdings>\[[^\]]*\])", re.S)
BUY_PATTERN = re.compile(r"今日买入信号[:：][ \t]*(?P<signals>[^\n]*)")
SELL_PATTERN = re.compile(r"今日卖出信号[:：][ \t]*(?P<signals>[^\n]*)")
DRAWDOWN_PATTERN = re.compile(r"当前回撤\s*(?P<current>[-+]?\d+(?:\.\d+)?)%，控仓法\s*(?P<control>[-+]?\d+(?:\.\d+)?)%")
ROW_PATTERN = re.compile(
    r"(?P<mark>[√×]{3})\s*(?P<name>.+?)\s*PB\s*(?P<pb>[0-9.]+)\s*(?P<suffix>[_A-Za-z0-9★]*)\s*:\s*\[(?P<values>.*?)\]",
    re.S,
)


def _json(value: object) -> str:
    """Serialize a value as compact UTF-8 JSON for CSV cells."""
    return json.dumps(value, ensure_ascii=False, separators=(",", ":"))


def _clean_text(text: str) -> str:
    """Normalize DOCX whitespace and signal fragments."""
    cleaned = text.replace("\xa0", " ").replace("\u3000", " ")
    cleaned = cleaned.replace("\r\n", "\n").replace("\r", "\n")
    cleaned = re.sub(r"(\d{4})\s*_\s*([BS])", r"\1_\2", cleaned)
    return cleaned


def read_docx_text(path: Path) -> str:
    """Read all paragraph and table text from a DOCX file."""
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover - optional dependency
        raise RuntimeError(f"python-docx is required to parse DOCX files: {exc}") from exc

    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return _clean_text("\n".join(parts))


def split_daily_sections(text: str) -> list[tuple[pd.Timestamp, str]]:
    """Split a multi-day DOCX text by ``今天是 YYYY-MM-DD`` sections."""
    cleaned = _clean_text(text)
    matches = list(DATE_PATTERN.finditer(cleaned))
    sections: list[tuple[pd.Timestamp, str]] = []
    for index, match in enumerate(matches):
        start = match.start()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(cleaned)
        sections.append((pd.Timestamp(match.group("date")).normalize(), cleaned[start:end]))
    return sections


def _extract_block(section: str, start_pattern: re.Pattern[str], end_pattern: re.Pattern[str] | None = None) -> str:
    """Extract a report block from a daily section."""
    start = start_pattern.search(section)
    if not start:
        return ""
    block_start = start.end()
    block_end = len(section)
    if end_pattern:
        end = end_pattern.search(section, block_start)
        if end:
            block_end = end.start()
    return section[block_start:block_end]


def _parse_holding_items(raw: str) -> tuple[list[dict[str, str]], list[str], list[str], list[str]]:
    """Parse a raw bracketed holding list into item metadata and clean names."""
    text = _clean_text(raw).strip()
    if not text:
        return [], [], [], []
    body = text.strip().strip("[]")
    if not body:
        return [], [], [], []
    parts = [part.strip().strip("'\"") for part in re.split(r"[,，]", body)]
    items: list[dict[str, str]] = []
    clean_names: list[str] = []
    added: list[str] = []
    removed: list[str] = []
    for part in parts:
        if not part:
            continue
        change = ""
        if "↑" in part:
            change = "added"
        elif "↓" in part:
            change = "removed"
        clean_name = normalize_name(part)
        item = {"raw": part, "clean_name": clean_name, "change": change}
        items.append(item)
        clean_names.append(clean_name)
        if change == "added":
            added.append(clean_name)
        elif change == "removed":
            removed.append(clean_name)
    return items, clean_names, added, removed


def parse_holdings(text: str) -> list[str]:
    """Parse a bracketed Chinese holding list and normalize names."""
    _, clean_names, _, _ = _parse_holding_items(text)
    return clean_names


def _dedupe(values: Iterable[str]) -> list[str]:
    """Return values in first-seen order without duplicates."""
    seen: set[str] = set()
    result: list[str] = []
    for value in values:
        if value and value not in seen:
            seen.add(value)
            result.append(value)
    return result


def parse_trade_signals(text: str) -> list[dict[str, str]]:
    """Parse buy/sell signal text into ``code``/``name`` dictionaries."""
    lines = _clean_text(text).splitlines()
    value = lines[0].strip() if lines else ""
    if not value or value in {"-", "无", "[]"}:
        return []
    signals: list[dict[str, str]] = []
    for match in re.finditer(r"(?P<code>\d{6}\.(?:SH|SZ))\s*\((?P<name>[^)]+)\)", value):
        signals.append({"code": match.group("code"), "name": normalize_name(match.group("name"))})
    if signals:
        return signals

    for part in re.split(r"[,，、;；]\s*", value):
        cleaned = part.strip().strip("'\"")
        if not cleaned:
            continue
        code_match = re.match(r"(?P<code>\d{6}\.(?:SH|SZ))$", cleaned)
        if code_match:
            signals.append({"code": code_match.group("code"), "name": ""})
        else:
            signals.append({"code": "", "name": normalize_name(cleaned)})
    return signals


def parse_signal_text(text: str) -> list[str]:
    """Parse buy/sell signal names from a report line."""
    return [item["name"] or item["code"] for item in parse_trade_signals(text)]


def _split_values(raw_values: str) -> list[str]:
    """Split a loose Python-list value string without trusting literal syntax."""
    text = _clean_text(raw_values)
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"(\d{4})\s*_\s*([BS])", r"\1_\2", text)
    values: list[str] = []
    current: list[str] = []
    quote: str | None = None
    for char in text:
        if char in {"'", '"'}:
            quote = None if quote == char else char if quote is None else quote
            current.append(char)
            continue
        if char in {",", "，"} and quote is None:
            values.append("".join(current).strip())
            current = []
        else:
            current.append(char)
    if current:
        values.append("".join(current).strip())
    return values


def _parse_scalar(value: str) -> object:
    """Parse one row value as float or normalized string."""
    text = value.strip().strip("'\"").strip()
    text = re.sub(r"\s+", "", text)
    text = re.sub(r"(\d{4})_?([BS])", r"\1_\2", text)
    if text in {"", "nan"}:
        return "-"
    try:
        return float(text)
    except ValueError:
        return text


def _float_or_none(value: object) -> float | None:
    """Convert a parsed scalar to float when possible."""
    try:
        if value == "-":
            return None
        return float(value)
    except (TypeError, ValueError):
        return None


def _literal_values(raw_values: str) -> list[object]:
    """Parse list-like row values into normalized scalars."""
    return [_parse_scalar(part) for part in _split_values(raw_values)]


def parse_rows(block: str, symbols: list[ETF], module: str = "broad") -> pd.DataFrame:
    """Parse single-ETF metric rows from one ETF block."""
    _, name_map = etf_maps(symbols)
    rows: list[dict[str, object]] = []
    for row_order, match in enumerate(ROW_PATTERN.finditer(_clean_text(block)), start=1):
        values = _literal_values(match.group("values"))
        values = values + ["-"] * (9 - len(values))
        name_raw = match.group("name").strip()
        name = normalize_name(name_raw)
        symbol = name_map.get(name)
        rows.append(
            {
                "row_order": row_order,
                "module": module,
                "mark": match.group("mark"),
                "name_raw": name_raw,
                "name": name,
                "code_if_found": symbol.code if symbol else "",
                "pb": _float_or_none(match.group("pb")),
                "suffix_flags": match.group("suffix") or "",
                "ret20_pct": _float_or_none(values[0]),
                "bias_pct": _float_or_none(values[1]),
                "trend": values[2],
                "market_signal": values[3],
                "convex": values[4],
                "nav_like": _float_or_none(values[5]),
                "last_signal": values[6],
                "profit_pct": _float_or_none(values[7]),
                "latest_signal": values[8],
                "raw_line": re.sub(r"\s+", " ", match.group(0)).strip(),
            }
        )
    return pd.DataFrame(rows)


def _parse_label(report_date: pd.Timestamp, broad_block: str) -> dict[str, object]:
    """Parse one daily broad-ETF label row."""
    broad_valid = bool(broad_block.strip()) and "等待数据更新" not in broad_block
    yesterday_match = YESTERDAY_PATTERN.search(broad_block)
    target_match = TARGET_PATTERN.search(broad_block)
    yesterday_raw = yesterday_match.group("holdings") if yesterday_match else ""
    target_raw = target_match.group("holdings") if target_match else ""

    yesterday_items, yesterday_holdings, _, removed_from_yesterday = _parse_holding_items(yesterday_raw)
    target_items, target_holdings, added_from_target, _ = _parse_holding_items(target_raw)
    added_names = _dedupe([*added_from_target, *[name for name in target_holdings if name not in yesterday_holdings]])
    removed_names = _dedupe(
        [*removed_from_yesterday, *[name for name in yesterday_holdings if name not in target_holdings]]
    )
    if not broad_valid:
        yesterday_holdings = []
        target_holdings = []
        added_names = []
        removed_names = []

    buy_match = BUY_PATTERN.search(broad_block)
    sell_match = SELL_PATTERN.search(broad_block)
    drawdown_match = DRAWDOWN_PATTERN.search(broad_block)
    position_state = "unknown"
    if "现在半仓" in broad_block:
        position_state = "half"
    elif "现在满仓" in broad_block:
        position_state = "full"

    return {
        "date": report_date.strftime("%Y-%m-%d"),
        "broad_valid": broad_valid,
        "yesterday_holdings_raw": yesterday_raw,
        "target_holdings_raw": target_raw,
        "yesterday_holdings": _json(yesterday_holdings),
        "target_holdings": _json(target_holdings),
        "yesterday_holding_items": _json(yesterday_items),
        "target_holding_items": _json(target_items),
        "added_names": _json(added_names),
        "removed_names": _json(removed_names),
        "add_position_signal": "可以加仓" in broad_block,
        "position_state": position_state,
        "current_drawdown_pct": float(drawdown_match.group("current")) if drawdown_match else None,
        "control_drawdown_pct": float(drawdown_match.group("control")) if drawdown_match else None,
        "buy_signals": _json(parse_trade_signals(buy_match.group("signals") if buy_match else "")),
        "sell_signals": _json(parse_trade_signals(sell_match.group("signals") if sell_match else "")),
    }


def parse_doc_text(text: str, symbols: list[ETF]) -> tuple[pd.DataFrame, pd.DataFrame]:
    """Parse multi-day report text into label and row tables."""
    sections = split_daily_sections(text)
    if not sections:
        raise ValueError("Could not parse any report date from DOCX text.")

    labels: list[dict[str, object]] = []
    row_frames: list[pd.DataFrame] = []
    for report_date, section in sections:
        broad_block = _extract_block(section, BROAD_START_PATTERN, SECTOR_START_PATTERN)
        sector_block = _extract_block(section, SECTOR_START_PATTERN)
        labels.append(_parse_label(report_date, broad_block))

        broad_rows = parse_rows(broad_block, symbols, module="broad")
        if not broad_rows.empty:
            broad_rows.insert(0, "date", report_date.strftime("%Y-%m-%d"))
            row_frames.append(broad_rows)

        sector_rows = parse_rows(sector_block, symbols, module="sector")
        if not sector_rows.empty:
            sector_rows.insert(0, "date", report_date.strftime("%Y-%m-%d"))
            row_frames.append(sector_rows)

    label_df = pd.DataFrame(labels).sort_values("date", ascending=False).reset_index(drop=True)
    rows_df = pd.concat(row_frames, ignore_index=True) if row_frames else pd.DataFrame()
    if not rows_df.empty:
        rows_df = rows_df.sort_values(["date", "module", "row_order"], ascending=[False, True, True]).reset_index(drop=True)
    return label_df, rows_df


def copy_default_doc_if_available(cfg: ProjectConfig) -> Path | None:
    """Copy ``/mnt/data/股票策略etf.docx`` into raw data when available."""
    source = Path("/mnt/data/股票策略etf.docx")
    target = cfg.raw_dir / "股票策略etf.docx"
    if source.exists() and not target.exists():
        ensure_dir(target.parent)
        shutil.copy2(source, target)
        return target
    return target if target.exists() else None


def parse_docx_file(path: Path, cfg: ProjectConfig) -> tuple[Path, Path]:
    """Parse a DOCX report and write label/row CSV files."""
    if not path.exists():
        default_doc = copy_default_doc_if_available(cfg)
        if default_doc and default_doc.exists():
            path = default_doc
        else:
            raise FileNotFoundError(f"DOCX file not found: {path}")

    symbols = load_universe(cfg.universe, pools=["broad_etf", "sector_etf"])
    text = read_docx_text(path)
    labels, rows = parse_doc_text(text, symbols)
    ensure_dir(cfg.labels_dir)
    label_path = cfg.labels_dir / "doc_labels.csv"
    row_path = cfg.labels_dir / "doc_rows.csv"
    labels.to_csv(label_path, index=False)
    rows.to_csv(row_path, index=False)
    return label_path, row_path
